import os
import docx
import openpyxl

# 获取当前目录下的所有 Word 文件
doc_files = [filename for filename in os.listdir('.') if filename.endswith('.docx')]

# 新建 Excel 文件
xlsx_file = '当前文件夹汇总.xlsx'
wb = openpyxl.Workbook()

# 获取 active sheet
ws = wb.active

# 设置第一行的列名
ws.cell(row=1, column=1).value = "zhCN"
ws.cell(row=1, column=2).value = "enUS"

# 从第二行开始逐行写入 Word 文档的内容
row_index = 2
for doc_file in doc_files:
    if not doc_file.startswith("~$"):
        # 打开 Word 文档
        doc = docx.Document(doc_file)

        # 判断第一段是否为一级标题，如果是则删除
        if doc.paragraphs[0].style.name == 'Heading 1':
            del doc.paragraphs[0]

        # 逐行写入中英文对照内容
        for p in range(len(doc.paragraphs)):
            if doc.paragraphs[p].text.strip(): # 确保段落非空
                if p % 2 == 0:
                    # 中文
                    ws.cell(row=row_index, column=1).value = doc.paragraphs[p].text
                else:
                    # 英文
                    ws.cell(row=row_index, column=2).value = doc.paragraphs[p].text
                    row_index += 1

# 删除空白行
for i in range(ws.max_row, 0, -1):
    if not ws.cell(row=i, column=1).value and not ws.cell(row=i, column=2).value:
        ws.delete_rows(i)

# 在汇总 Excel 的首行第一列和第二列重新标注 zhCN 和 enUS
ws.cell(row=1, column=1).value = "zhCN"
ws.cell(row=1, column=2).value = "enUS"

# 保存 Excel 文件
wb.save(xlsx_file)
