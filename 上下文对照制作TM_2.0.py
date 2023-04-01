import os
import docx
import openpyxl

# 获取当前目录下的所有 Word 文件
doc_files = [filename for filename in os.listdir('.') if filename.endswith('.docx')]

for doc_file in doc_files:
    # 打开 Word 文档
    doc = docx.Document(doc_file)

    # 新建 Excel 文件
    xlsx_file = os.path.splitext(doc_file)[0] + '.xlsx'
    wb = openpyxl.Workbook()

    # 获取 active sheet
    ws = wb.active

    # 设置第一行的列名
    ws.cell(row=1, column=1).value = "zhCN"
    ws.cell(row=1, column=2).value = "enUS"

    # 提取中英文对照内容并写入 Excel
    for i in range(len(doc.paragraphs)):
        if i % 2 == 0:
            # 中文
            ws.cell(row=i // 2 + 2, column=1).value = doc.paragraphs[i].text
        else:
            # 英文
            ws.cell(row=i // 2 + 2, column=2).value = doc.paragraphs[i].text

    # 保存 Excel 文件
    wb.save(xlsx_file)
